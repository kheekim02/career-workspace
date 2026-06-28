#!/usr/bin/env python3
"""
Generate an ATS-optimized DOCX resume for Kunhee (Geoffrey) Kim.

Design choices follow 2026 ATS best practices:
- Single-column, no tables/text boxes/columns (parsers read linear top-to-bottom).
- Standard web-safe font (Calibri), 10.5pt body, larger name/headers.
- Contact info in the document body (not Word header/footer, which parsers skip).
- Standard section headings (Experience, Projects, Technical Skills, Education).
- Consistent Month YYYY dates; plain bullets; skills echoed in experience bullets.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
FONT = "Calibri"

# Set PHONE to your number (e.g. "(510) 555-0123") to include it in the contact line.
PHONE = ""


def set_margins(doc, inches=0.6):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.05


def add_run(p, text, size=10.5, bold=False, italic=False, color=INK):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    return r


def name_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "Kunhee (Geoffrey) Kim", size=20, bold=True, color=ACCENT)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    add_run(p2, "DATA ENGINEER  |  DATA ANALYST", size=10.5, bold=True, color=MUTED)

    contact_bits = ["kunheekim02@gmail.com"]
    if PHONE:
        contact_bits.append(PHONE)
    contact_bits += ["Berkeley / Bay Area, CA", "linkedin.com/in/kk02", "github.com/kheekim02"]
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(4)
    add_run(p3, "  |  ".join(contact_bits), size=9.5, color=MUTED)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text.upper(), size=11, bold=True, color=ACCENT)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pbdr.append(bottom)
    pPr.append(pbdr)


def body_para(doc, text, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    add_run(p, text)
    return p


def entry_header(doc, left, right=None, left_bold=True):
    """Title left, date right, on one line using a right-aligned tab stop (not a table)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, left, bold=left_bold)
    if right:
        # right tab at the usable width
        usable = Inches(8.5 - 2 * 0.6)
        tab_pos = int(usable)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(tab_pos))
        tabs.append(tab)
        pPr.append(tabs)
        add_run(p, "\t")
        add_run(p, right, size=9.5, italic=True, color=MUTED)
    return p


def tech_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, text, size=9.5, italic=True, color=MUTED)


def bullet(doc, runs):
    """runs: list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.05
    for text, bold in runs:
        add_run(p, text, bold=bold)


def skill(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    add_run(p, label + ": ", bold=True, color=ACCENT)
    add_run(p, value)


def build():
    doc = Document()
    set_margins(doc, 0.6)
    style_base(doc)

    name_block(doc)

    heading(doc, "Summary")
    body_para(doc,
        "Data Engineer and UC Berkeley Data Science graduate who designs and scales production "
        "ETL/ELT pipelines, multi-terabyte data lakes and warehouses (PostgreSQL, Neo4j, Apache "
        "Spark), and ML/NLP systems. Scaled a research data lake to 2M+ filings and 33M+ records "
        "with strict data-integrity guarantees, relational data modeling, and SQL performance "
        "tuning. Bilingual in English and Korean.", after=2)

    heading(doc, "Experience")
    entry_header(doc, "Financial Data Analyst — Global Key Advisors, Bay Area, CA",
                 "January 2026 – Present")
    bullet(doc, [("Re-architected SEC EDGAR extract-transform-load (ETL) scraping pipelines "
                  "(8-K, 10-K, 10-Q, Form 15/25) from a memory-constrained CSV approach to a threaded "
                  "PostgreSQL integration, scaling the data lake to ", False),
                 ("2.05M+ text filings and 3.3M+ financial records across 28,000+ entities", True),
                 (".", False)])
    bullet(doc, [("Administered a PostgreSQL research warehouse (sec_data) holding ", False),
                 ("33M+ records and 42 GB of text", True),
                 ("; architected modular schemas, tuned TOAST storage, and debugged the query planner "
                  "via system catalogs and ANALYZE statistics.", False)])
    bullet(doc, [("Automated daily ingestion of a 1.12 GB price panel from a remote Windows host to a "
                  "Linux Spark cluster, ", False),
                 ("eliminating IP-ban risk and saving hours/week", True),
                 ("; built data-integrity logic (CIK joins, point-in-time mapping, vendor dedup, "
                  "idempotent upserts) to prevent silent corruption.", False)])
    bullet(doc, [("Built NLP extraction with local LLMs (Ollama) parsing filings into "
                  "corporate-relationship edges in a ", False),
                 ("Neo4j knowledge graph", True),
                 (" via an async pipeline processing ", False),
                 ("~480 filings/hour", True),
                 (" unattended.", False)])
    bullet(doc, [("Migrated prediction strategy from FinBERT sentiment to Graph Neural Networks (GNN); "
                  "out-of-sample backtesting ", False),
                 ("exposed lookahead bias", True),
                 (" (a perceived +2.39% return collapsed to negative under 30-bps costs), preventing a "
                  "losing strategy deployment.", False)])

    heading(doc, "Projects")
    entry_header(doc, "Alpha Signals — Quantitative SEC Event Research Pipeline",
                 "Jan 2026 – Present")
    tech_line(doc, "Python, Polars, PostgreSQL, Neo4j, Ollama, Graph Neural Networks")
    bullet(doc, [("Flagship pipeline testing whether corporate events in SEC filings predict stock "
                  "returns; built an autonomous daily loop (signal generation to simulated trades to GNN "
                  "predictions by 04:30 PDT) over a Neo4j knowledge graph spanning 28,000+ entities.", False)])
    entry_header(doc, "PathWise — Intelligent iOS Running App", "Jan 2026 – May 2026")
    tech_line(doc, "Swift (SwiftUI, MapKit, ActivityKit), FastAPI, PostgreSQL + PostGIS, GraphHopper, Docker")
    bullet(doc, [("Full-stack architect: built a pathfinding orchestrator integrating GraphHopper with "
                  "PostGIS for real-time geofenced routing and a native iOS client with turn-by-turn voice "
                  "navigation; achieved ", False),
                 ("sub-2-second route generation", True),
                 (". Presented at Berkeley's startup pitch competition.", False)])

    heading(doc, "Technical Skills")
    skill(doc, "Programming Languages", "Python, SQL, Swift, Bash/Shell")
    skill(doc, "Data Engineering", "ETL/ELT pipeline design, data lake & warehouse architecture, "
          "relational data modeling, schema design, idempotent ingestion, data quality & validation, "
          "deduplication, web scraping")
    skill(doc, "Databases & Big Data", "PostgreSQL (performance tuning, query optimization), Neo4j, "
          "Apache Spark (PySpark), PostGIS, SQLite")
    skill(doc, "Python & Backend", "Pandas, Polars, FastAPI, SQLAlchemy, Pydantic, concurrent.futures, "
          "BeautifulSoup, REST APIs")
    skill(doc, "ML / NLP", "Local LLM inference (Ollama), Graph Neural Networks (GNN), entity "
          "resolution, backtesting, event studies")
    skill(doc, "Tools & Workflow", "Docker, Linux, Git, cron/systemd orchestration, "
          "pg_dump/pg_restore, Tableau")
    skill(doc, "Languages (Spoken)", "English (fluent), Korean (fluent)")

    heading(doc, "Education")
    entry_header(doc, "University of California, Berkeley — B.A. Data Science", "May 2026")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "GPA 3.81 | Completed in 3 years | 2-year military service in Korea (2022–2024)",
            size=9.5, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, "Certificates: Berkeley Changemaker; Entrepreneurship & Technology (Sutardja Center / "
            "SCET, UC Berkeley)")

    heading(doc, "Research & Leadership")
    bullet(doc, [("Student Researcher, UC Berkeley", True),
                 (" (Prof. Jill Berrick, FamiliesNotFees) — Applied statistical modeling and ML to "
                  "national child-welfare data; built Tableau dashboards and the project website, growing "
                  "reach to 4,000+ views / 1,000+ unique visitors.", False)])
    bullet(doc, [("Course Coordinator, ENGIN 270B", True),
                 (" (2025) — Ran operations for a graduate course of 100+ students; built automated Excel "
                  "tracking surfacing data-driven metrics for curriculum improvement.", False)])
    bullet(doc, [("Berkeley Data Discovery", True),
                 (" — Meta-analysis of national transfer metrics using ML on large-scale education "
                  "datasets, with interactive dashboards for strategic decision-making.", False)])

    out = "Kunhee-Kim-Data-Engineer.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build()
